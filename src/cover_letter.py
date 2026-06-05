"""Cover letter docx builder — simple, ATS-safe.

If the body starts with the FAILURE_SENTINEL string, render a fail-loud banner
instead of a normal letter so the user can never accidentally submit a letter
where AI generation broke. The banner runs in red and the sign-off is omitted.
"""
from __future__ import annotations
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

# Imported here to keep cover_letter.py decoupled from the LLM pipeline; the
# string itself is the contract between the two modules.
FAILURE_SENTINEL = "[GENERATION FAILED — review before sending]"


def build_cover_letter(
    body_text: str,
    user: dict,
    job: dict,
    out_path: Path,
    *,
    font: str = "Calibri",
    base_size: float = 11.0,
    margins: float = 1.0,
):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(base_size)

    for section in doc.sections:
        section.top_margin = Inches(margins)
        section.bottom_margin = Inches(margins)
        section.left_margin = Inches(margins)
        section.right_margin = Inches(margins)

    def _add(text, *, bold=False, after=4, size=None, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size or base_size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return p

    is_failure = body_text.strip().startswith(FAILURE_SENTINEL)
    body_for_render = body_text
    if is_failure:
        body_for_render = body_text.strip()[len(FAILURE_SENTINEL):].lstrip("\n").strip()

    # Header block
    _add(user["full_name"], bold=True, after=0)
    _add(user.get("email", ""), after=0, size=base_size - 0.5)
    if user.get("phone"):
        _add(user["phone"], after=0, size=base_size - 0.5)
    if user.get("location_city"):
        _add(user["location_city"], after=8, size=base_size - 0.5)

    _add(date.today().strftime("%B %d, %Y"), after=8)

    _add("Hiring Team", after=0)
    _add(job["company"], after=0, bold=True)
    if job.get("location"):
        _add(job["location"], after=10, size=base_size - 0.5)

    # Skip the auto-salutation if the model body already starts with one,
    # so future prompt tweaks that allow "Dear ..." don't double up.
    if not is_failure and not body_for_render.lstrip().lower().startswith("dear"):
        _add("Dear Hiring Team,", after=8)

    if is_failure:
        # Loud red banner. Two lines so it survives even rough ATS extraction.
        red = RGBColor(0xC0, 0x00, 0x00)
        _add(
            "AUTOMATED GENERATION FAILED -- DO NOT SEND",
            bold=True, after=2, size=base_size + 1, color=red,
        )
        _add(
            "The AI pipeline could not produce a clean cover letter for this role. "
            "The text below is the best raw recovery; rewrite it manually before "
            "submitting. Use 'python -m src.tweak <folder>' to regenerate, or copy "
            "this draft into Word and revise.",
            after=10, size=base_size - 0.5, color=red,
        )
        # Render whatever recovery text we have, clearly de-emphasized.
        if body_for_render:
            for para in body_for_render.split("\n\n"):
                if para.strip():
                    _add(para.strip(), after=8)
        else:
            _add("(no recoverable content from the model)", after=8,
                 size=base_size - 0.5)
        # Skip sign-off — this is a draft, not a finished letter.
    else:
        # Normal cover letter render.
        for para in body_for_render.strip().split("\n\n"):
            _add(para.strip(), after=8)

        _add("Best,", after=0)
        _add(user["full_name"], bold=True, after=0)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
