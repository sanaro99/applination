"""
Build an ATS-safe .docx resume from the structured JSON the tailor produces.

Style reference: a clean one-page ATS resume.
  Font: Times New Roman, body 9pt, headings 11pt bold
  Margins: 0.19" top/bottom, 0.25" left/right
  Spacing: space_before=2pt, space_after=0, line_spacing=1.0

One-page enforcement:
- The tailor already constrains content to measured budgets.
- After write, we estimate page count and, if >1, iteratively drop lowest-priority
  content (coursework, 3rd project, extra bullets, etc.) and retry.
"""
from __future__ import annotations
from copy import deepcopy
from pathlib import Path
import logging

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOG = logging.getLogger(__name__)


# ---------------------------------------------------------------------
# Skills normalization — handle various model output shapes
# ---------------------------------------------------------------------

# Friendly display names for group keys the LLM (or master YAML) might emit
# in snake_case or other ATS-unfriendly forms.
_GROUP_NAME_MAP = {
    "languages": "Languages",
    "ai_ml": "AI / ML",
    "ai/ml": "AI / ML",
    "ai": "AI / ML",
    "ml": "AI / ML",
    "machine_learning": "AI / ML",
    "web_and_apis": "Frameworks & APIs",
    "web": "Frameworks & APIs",
    "frameworks": "Frameworks & APIs",
    "frameworks_and_libraries": "Frameworks & Libraries",
    "devops_cloud": "Cloud & DevOps",
    "cloud": "Cloud & DevOps",
    "devops": "Cloud & DevOps",
    "sre_infra": "Infrastructure & SRE",
    "sre": "Infrastructure & SRE",
    "infra": "Infrastructure & SRE",
    "infrastructure": "Infrastructure & SRE",
    "data": "Data & Storage",
    "databases": "Data & Storage",
    "data_storage": "Data & Storage",
    "tools": "Tools",
    "practices": "Engineering Practices",
    "engineering_practices": "Engineering Practices",
    "soft_skills": "Engineering Practices",
}


def _prettify_group(name: str) -> str:
    """Map snake_case / lowercase group keys to ATS-friendly display names."""
    if not name:
        return "Skills"
    key = name.strip().lower().replace(" ", "_").replace("-", "_").replace("/", "_")
    if key in _GROUP_NAME_MAP:
        return _GROUP_NAME_MAP[key]
    # Already cased nicely? Keep as-is. Otherwise title-case the words.
    if any(c.isupper() for c in name) and "_" not in name:
        return name.strip()
    return " ".join(w.capitalize() for w in name.replace("_", " ").split())


def _normalize_skills(skills) -> list[dict]:
    """Normalize skills to list[{'group': str, 'items': list[str]}].

    Handles all shapes that models (and the master YAML) emit:
      - Correct:  [{'group': 'Languages', 'items': ['Python', 'SQL']}, ...]
      - Dict:     {'languages': ['Python', 'SQL'], 'ai_ml': [...]}        # YAML / loose LLM
      - Flat str: ['Python', 'SQL', 'Docker', ...]
      - Colon str:['Languages: Python, SQL', 'Cloud: AWS, GCP', ...]
    """
    if not skills:
        return []

    normalized: list[dict] = []

    # --- Dict shape: {'group_key': [items...]} -----------------------------
    # This is what the LLM returns most often when it ignores the schema —
    # previously we silently dropped every item.
    if isinstance(skills, dict):
        for group_key, items in skills.items():
            if isinstance(items, str):
                items = [x.strip() for x in items.split(",") if x.strip()]
            elif not isinstance(items, list):
                continue
            normalized.append({
                "group": _prettify_group(str(group_key)),
                "items": [str(i).strip() for i in items if str(i).strip()],
            })
        return [g for g in normalized if g["items"]]

    # --- List shape (correct or strings) -----------------------------------
    for item in skills:
        if isinstance(item, dict):
            items = item.get("items", [])
            if isinstance(items, str):
                items = [x.strip() for x in items.split(",") if x.strip()]
            normalized.append({
                "group": _prettify_group(str(item.get("group", "Skills"))),
                "items": [str(i).strip() for i in items if str(i).strip()],
            })
        elif isinstance(item, str):
            if ": " in item:
                group, rest = item.split(": ", 1)
                normalized.append({
                    "group": _prettify_group(group.strip()),
                    "items": [x.strip() for x in rest.split(",") if x.strip()],
                })
            else:
                if normalized and normalized[-1]["group"] == "Skills":
                    normalized[-1]["items"].append(item.strip())
                else:
                    normalized.append({"group": "Skills", "items": [item.strip()]})

    # Dedupe items within each group (case-insensitive, preserving order)
    for g in normalized:
        seen = set()
        unique = []
        for it in g["items"]:
            k = it.lower()
            if k not in seen:
                seen.add(k)
                unique.append(it)
        g["items"] = unique

    return [g for g in normalized if g["items"]]


# ---------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------
def _set_margins(doc: Document, left_right: float = 0.25, top_bottom: float = 0.19):
    for section in doc.sections:
        section.top_margin = Inches(top_bottom)
        section.bottom_margin = Inches(top_bottom)
        section.left_margin = Inches(left_right)
        section.right_margin = Inches(left_right)


def _style_run(run, *, size: float, bold: bool = False, font: str = "Times New Roman",
               color: tuple[int, int, int] = (20, 20, 20)):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def _tight_para(p, space_before: float = 2, space_after: float = 0, line: float = 1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line


def _add_heading(doc: Document, text: str, *, font: str, size: float,
                 space_before: float = 7):
    """Bold section heading with a thin bottom rule — standard professional style."""
    p = doc.add_paragraph()
    _tight_para(p, space_before=space_before, space_after=2)
    r = p.add_run(text.upper())
    _style_run(r, size=size, bold=True, font=font, color=(20, 20, 20))

    # Thin bottom border under the heading (gives the classic "SECTION _______" look)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")        # 0.5pt line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "404040")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_bullet(doc: Document, text: str, *, font: str, size: float):
    p = doc.add_paragraph(style=None)
    _tight_para(p, space_before=1, space_after=0, line=1.0)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.18)
    r = p.add_run("• ")
    _style_run(r, size=size, font=font)
    r2 = p.add_run(text)
    _style_run(r2, size=size, font=font)
    return p


def _add_bullet_split(doc: Document, bold_prefix: str, plain_suffix: str, *,
                      font: str, size: float):
    """Bullet line with a bolded leading run + regular trailing run.

    Used for awards: the name is bolded for recruiter-scan + ATS keyword
    weight, the description and date follow in regular weight.
    """
    p = doc.add_paragraph(style=None)
    _tight_para(p, space_before=1, space_after=0, line=1.0)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.18)
    r0 = p.add_run("• ")
    _style_run(r0, size=size, font=font)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        _style_run(r1, size=size, bold=True, font=font)
    if plain_suffix:
        r2 = p.add_run(plain_suffix)
        _style_run(r2, size=size, bold=False, font=font)
    return p


def _hex(rgb: tuple[int, int, int]) -> str:
    return "%02X%02X%02X" % rgb


def _add_hyperlink(paragraph, url: str, text: str, *, font: str, size: float,
                   bold: bool = False,
                   color: tuple[int, int, int] = (20, 20, 20)):
    """Insert a clickable hyperlink run into paragraph. Underlined, no default blue."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    href = url if url.startswith(("http://", "https://", "mailto:", "tel:")) \
        else f"https://{url}"
    r_id = paragraph.part.relate_to(href, RT.HYPERLINK, is_external=True)

    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rPr.append(rFonts)

    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(int(size * 2)))
        rPr.append(el)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), _hex(color))
    rPr.append(color_el)

    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)

    if bold:
        rPr.append(OxmlElement("w:b"))

    run_el.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run_el.append(t)

    hl.append(run_el)
    paragraph._p.append(hl)
    return hl


def _contact_href(kind: str, val: str) -> str | None:
    """Return the href for a contact item, or None for plain-text items."""
    if kind == "email":
        return f"mailto:{val}"
    if kind == "phone":
        digits = "".join(c for c in val if c.isdigit() or c == "+")
        return f"tel:{digits}" if digits else None
    if kind in ("linkedin", "github", "url"):
        return val if val.startswith(("http://", "https://")) else f"https://{val}"
    return None


# ---------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------
def build_resume_docx(
    resume: dict,
    user: dict,
    out_path: Path,
    *,
    font: str = "Times New Roman",
    base_size: float = 9.0,
    margins: float = 0.25,
):
    resume = deepcopy(resume)
    resume["skills"] = _normalize_skills(resume.get("skills", []))

    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(base_size)

    # Margins: left/right = margins, top/bottom proportionally tighter
    top_bottom = max(0.19, margins * 0.76)
    _set_margins(doc, left_right=margins, top_bottom=top_bottom)

    heading_size = base_size + 2  # 9 + 2 = 11pt for headings

    # -------- NAME + CONTACT --------
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_para(name_p, space_before=0, space_after=0)
    nr = name_p.add_run(user["full_name"])
    _style_run(nr, size=base_size + 13, bold=False, font=font, color=(10, 10, 10))

    contact_bits = [
        ("email",    user.get("email", "")),
        ("phone",    user.get("phone", "")),
        ("linkedin", user.get("linkedin", "")),
        ("github",   user.get("github", "")),
        ("url",      user.get("portfolio", "")),
        ("text",     user.get("location_city", "")),
    ]
    contact_bits = [(kind, v) for kind, v in contact_bits if v]

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight_para(contact_p, space_before=2, space_after=0)

    for i, (kind, val) in enumerate(contact_bits):
        if i > 0:
            sep = contact_p.add_run("   |   ")
            _style_run(sep, size=base_size - 0.5, font=font, color=(130, 130, 130))

        href = _contact_href(kind, val)
        if href:
            _add_hyperlink(contact_p, href, val, font=font, size=base_size,
                           color=(60, 60, 60))
        else:
            r = contact_p.add_run(val)
            _style_run(r, size=base_size, font=font, color=(60, 60, 60))

    # -------- SUMMARY --------
    if resume.get("summary"):
        _add_heading(doc, "Professional Summary", font=font, size=heading_size, space_before=7)
        p = doc.add_paragraph()
        _tight_para(p, space_before=2, space_after=0, line=1.0)
        r = p.add_run(resume["summary"])
        _style_run(r, size=base_size, font=font)

    # -------- EDUCATION --------
    if resume.get("education"):
        _add_heading(doc, "Education", font=font, size=heading_size, space_before=7)
        for edu in resume["education"]:
            # Line 1: School, Location  ...  Dates
            p = doc.add_paragraph()
            _tight_para(p, space_before=2, space_after=0)
            # Drop location tokens already present in the school name so we
            # don't render "University of Washington, Seattle, Seattle, WA".
            # Word-boundary match so "WA" isn't dropped just because
            # "wa" is a substring of "washington".
            school = edu["school"]
            loc_raw = edu.get("location", "")
            import re as _re
            school_words = {w for w in _re.findall(r"[A-Za-z0-9]+", school.lower())}
            loc_tokens = [t.strip() for t in loc_raw.split(",")
                          if t.strip() and t.strip().lower() not in school_words]
            loc = ", ".join(loc_tokens)
            r1 = p.add_run(f"{school}, {loc}" if loc else school)
            _style_run(r1, size=base_size, bold=True, font=font)
            tail = f"   {edu.get('dates', '')}"
            r2 = p.add_run(tail)
            _style_run(r2, size=base_size, bold=True, font=font, color=(80, 80, 80))

            # Line 2: Degree, Minor, GPA
            p2 = doc.add_paragraph()
            _tight_para(p2, space_before=0, space_after=0)
            deg_text = edu["degree"]
            if edu.get("minor"):
                deg_text += f", Minor in {edu['minor']}"
            if edu.get("gpa"):
                deg_text += f"; GPA: {edu['gpa']}"
            r3 = p2.add_run(deg_text)
            _style_run(r3, size=base_size, font=font)

            # Line 2b: Specializations (some degrees have these — were silently dropped)
            specs = edu.get("specializations")
            if specs:
                if isinstance(specs, list):
                    specs_text = ", ".join(str(s).strip() for s in specs if s)
                else:
                    specs_text = str(specs)
                if specs_text:
                    p_sp = doc.add_paragraph()
                    _tight_para(p_sp, space_before=0, space_after=0)
                    r_sp = p_sp.add_run(f"Specializations: {specs_text}")
                    _style_run(r_sp, size=base_size, font=font, color=(60, 60, 60))

            # Line 2c: Honors (Dean's List, etc.) — italic-feel via grey color
            honors = edu.get("honors")
            if honors:
                if isinstance(honors, list):
                    honors_text = ", ".join(str(h).strip() for h in honors if h)
                else:
                    honors_text = str(honors)
                if honors_text:
                    p_h = doc.add_paragraph()
                    _tight_para(p_h, space_before=0, space_after=0)
                    r_h = p_h.add_run(f"Honors: {honors_text}")
                    _style_run(r_h, size=base_size, font=font, color=(60, 60, 60))

            # Line 3: Coursework
            cw = edu.get("coursework")
            if cw:
                if isinstance(cw, list):
                    cw = ", ".join(str(x) for x in cw)
                p3 = doc.add_paragraph()
                _tight_para(p3, space_before=0, space_after=0)
                rc = p3.add_run(f"Coursework: {cw}")
                _style_run(rc, size=base_size, font=font, color=(60, 60, 60))

    # -------- SKILLS --------
    if resume.get("skills"):
        _add_heading(doc, "Skills", font=font, size=heading_size, space_before=7)
        for group in resume["skills"]:
            p = doc.add_paragraph()
            _tight_para(p, space_before=2, space_after=0, line=1.0)
            # Group label bold; items normal weight so they don't compete visually.
            r_label = p.add_run(f"{group['group']}: ")
            _style_run(r_label, size=base_size, bold=True, font=font)
            r_items = p.add_run(", ".join(group["items"]))
            _style_run(r_items, size=base_size, bold=False, font=font)

    # -------- EXPERIENCE --------
    if resume.get("experience"):
        _add_heading(doc, "Experience", font=font, size=heading_size, space_before=7)
        for exp in resume["experience"]:
            p = doc.add_paragraph()
            _tight_para(p, space_before=2, space_after=0)
            role_co = f"{exp['role']} | {exp['company']}"
            if exp.get("location"):
                role_co += f", {exp['location']}"
            dates = exp.get("dates", "")
            r1 = p.add_run(role_co)
            _style_run(r1, size=base_size, bold=True, font=font)
            if dates:
                r2 = p.add_run(f"   {dates}")
                _style_run(r2, size=base_size, bold=True, font=font, color=(80, 80, 80))
            for b in exp.get("bullets", []):
                _add_bullet(doc, b, font=font, size=base_size)

    # -------- PROJECTS --------
    if resume.get("projects"):
        _add_heading(doc, "Projects", font=font, size=heading_size, space_before=7)
        for proj in resume["projects"]:
            p = doc.add_paragraph()
            _tight_para(p, space_before=2, space_after=0)
            r1 = p.add_run(proj["name"])
            _style_run(r1, size=base_size, bold=True, font=font)
            if proj.get("link"):
                sep = p.add_run("  |  ")
                _style_run(sep, size=base_size, bold=False, font=font, color=(100, 100, 100))
                _add_hyperlink(p, proj["link"], proj["link"],
                               font=font, size=base_size, bold=False)
            tech = proj.get("tech", "")
            if tech:
                r2 = p.add_run(f"   {tech}")
                _style_run(r2, size=base_size, font=font, color=(80, 80, 80))
            for b in proj.get("bullets", []):
                _add_bullet(doc, b, font=font, size=base_size)

    # -------- CERTIFICATIONS (one per line — ATS parsers split reliably) --------
    if resume.get("certifications"):
        certs_raw = resume["certifications"]
        certs = certs_raw if isinstance(certs_raw, list) else [certs_raw]
        certs = [str(c).strip() for c in certs if c and str(c).strip()]
        if certs:
            _add_heading(doc, "Certifications", font=font, size=heading_size, space_before=7)
            for c in certs:
                _add_bullet(doc, c, font=font, size=base_size)

    # -------- AWARDS & HONORS (one per line, bolded name) --------
    if resume.get("awards"):
        awards_raw = resume["awards"]
        awards = awards_raw if isinstance(awards_raw, list) else [awards_raw]
        if awards:
            _add_heading(doc, "Awards & Honors", font=font, size=heading_size, space_before=7)
            for a in awards:
                if isinstance(a, dict):
                    name = (a.get("name") or "").strip()
                    desc = (a.get("description") or "").strip()
                    date_s = (a.get("date") or "").strip()
                    if not name:
                        continue
                    suffix_parts = []
                    if desc:
                        suffix_parts.append(f", {desc}")
                    if date_s:
                        suffix_parts.append(f" ({date_s})")
                    _add_bullet_split(
                        doc, name, "".join(suffix_parts),
                        font=font, size=base_size,
                    )
                elif a:
                    _add_bullet(doc, str(a).strip(), font=font, size=base_size)

    # -------- ACTIVITIES & LEADERSHIP (one per line, optional) --------
    if resume.get("activities"):
        acts_raw = resume["activities"]
        acts = acts_raw if isinstance(acts_raw, list) else [acts_raw]
        acts = [str(a).strip() for a in acts if a and str(a).strip()]
        if acts:
            _add_heading(doc, "Activities & Leadership",
                         font=font, size=heading_size, space_before=7)
            for a in acts:
                _add_bullet(doc, a, font=font, size=base_size)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    LOG.debug("wrote %s", out_path)


# ---------------------------------------------------------------------
# One-page enforcement — both shrink overflow AND expand undersized
# ---------------------------------------------------------------------
def _estimate_line_count(resume: dict, user: dict | None = None,
                         base_size: float = 9.0) -> int:
    """Estimate body-line equivalents for the given body font size.

    Body line  = base_size + 2pt space_before
    Heading    = (base_size+2)pt font + 7pt space_before → counts as 1.6 body lines
    Name row   = large; contact row                       → counts as 3 lines combined

    CPL (chars/printed line) is font-aware — see _cpl(). The previous fixed
    CPL=108 assumed 9pt but under-counted real capacity, so a ~125-char bullet
    that fits one line at 10pt was counted as two lines, inflating the estimate
    and dropping content that actually fit. See _page_budget() for the matching
    vertical budget.
    """
    CPL = _cpl(base_size)   # chars per printed line at this font size
    HDR = 1.6               # heading weight relative to one body line

    lines: float = 0
    lines += 3.0  # name (large) + contact line

    if resume.get("summary"):
        lines += HDR  # heading
        lines += max(1, (len(resume["summary"]) + CPL - 1) // CPL)

    if resume.get("education"):
        lines += HDR  # heading
        for e in resume["education"]:
            lines += 2  # school + degree line
            if e.get("specializations"):
                lines += 1
            if e.get("honors"):
                lines += 1
            if e.get("coursework"):
                lines += 1

    skills = _normalize_skills(resume.get("skills", []))
    if skills:
        lines += HDR  # heading
        for g in skills:
            txt = f"{g['group']}: {', '.join(g['items'])}"
            lines += max(1, (len(txt) + CPL - 1) // CPL)

    if resume.get("experience"):
        lines += HDR  # heading
        for e in resume["experience"]:
            lines += 1  # role/company header
            for b in e.get("bullets", []):
                lines += max(1, (len(b) + CPL - 1) // CPL)

    if resume.get("projects"):
        lines += HDR  # heading
        for p in resume["projects"]:
            lines += 1  # project name/link header
            for b in p.get("bullets", []):
                lines += max(1, (len(b) + CPL - 1) // CPL)

    if resume.get("certifications"):
        cert_count = len(resume["certifications"]) if isinstance(resume["certifications"], list) else 1
        lines += HDR + cert_count  # heading + one row per cert

    if resume.get("awards"):
        award_count = len(resume["awards"]) if isinstance(resume["awards"], list) else 1
        lines += HDR + award_count  # heading + one row per award

    if resume.get("activities"):
        act_count = len(resume["activities"]) if isinstance(resume["activities"], list) else 1
        lines += HDR + act_count  # heading + one row per activity

    return int(lines)


# Page geometry — font-aware. Calibrated against real Word page counts at the
# 10pt project default (review 2026-05-24): a budget of 62 body-line-equivalents
# cleanly separates 1-page resumes (est <= 62) from 2-page ones (est ~66). Both
# the budget and the per-line char capacity scale with the body font size, so
# the page fit stays correct if output.base_font_size changes.
_CAL_FONT = 10.0          # font the calibration below was measured at
# Measured Word page counts at 10pt (review 2026-05-24): est 59 -> 1 page,
# est 61 -> 2 pages. Budget 60 is the safe cutoff (shrink anything above; never
# trims a legit 1-page resume that estimates <=59). Minimum 54 keeps a tidy
# single-line-bullet resume (which estimates low precisely because it is space-
# efficient) from triggering aggressive back-fill that overshoots to 2 pages.
_CAL_BUDGET = 60          # body-line-equivalents that fit one page at _CAL_FONT
# A space-efficient resume (single-line bullets) estimates LOW even when it is a
# full, content-rich page: measured, est 47 renders ~89% of page height (nearly
# full). So expansion fires only when GENUINELY sparse (est < 45) — otherwise
# adding a whole awards/certs block overshoots the little remaining room and
# spills to a 2nd page.
_CAL_MINIMUM = 45         # below this a page looks visibly half-empty
_CAL_CHARS_PER_LINE_AT_1PT = 1320.0   # CPL = round(this / font_pt); 10pt -> 132


def _cpl(base_size: float) -> int:
    """Characters per printed bullet line at the given body font size."""
    return round(_CAL_CHARS_PER_LINE_AT_1PT / (float(base_size) or _CAL_FONT))


def _page_budget(base_size: float) -> int:
    """Body-line-equivalents that fit one page (taller font -> fewer lines)."""
    return round(_CAL_BUDGET * (_CAL_FONT + 2) / ((float(base_size) or _CAL_FONT) + 2))


def _page_minimum(base_size: float) -> int:
    return round(_CAL_MINIMUM * (_CAL_FONT + 2) / ((float(base_size) or _CAL_FONT) + 2))


# Back-compat module constants (10pt values) for any external readers.
PAGE_LINE_BUDGET = _page_budget(_CAL_FONT)
PAGE_LINE_MINIMUM = _page_minimum(_CAL_FONT)


# ---------------------------------------------------------------------
# Shrink helpers (used when the resume overflows the page)
# ---------------------------------------------------------------------
def _drop_coursework(r: dict) -> dict:
    for e in r.get("education", []):
        e.pop("coursework", None)
    return r


def _drop_last(r: dict, key: str) -> dict:
    items = r.get(key) or []
    if items:
        items.pop()
    return r


def _truncate_bullets(r: dict, key: str, max_bullets: int) -> dict:
    for item in r.get(key, []) or []:
        item["bullets"] = (item.get("bullets") or [])[:max_bullets]
    return r


def _keep_top_projects(r: dict, keep: int) -> dict:
    """Trim the projects list to at most `keep`, preserving order (best first)."""
    projs = r.get("projects") or []
    if len(projs) > keep:
        r["projects"] = projs[:keep]
    return r


def _shrink(r: dict, base_size: float = _CAL_FONT) -> dict:
    """Iteratively drop lowest-priority content until estimate <= budget.

    Order matters: shed the least valuable content first. We protect a floor of
    TWO projects — trimming experience/project bullet COUNTS and dropping
    coursework/certs before ever dropping a project below two, since a resume
    with a single project reads thin. Only as a last resort do we fall to one
    project.
    """
    budget = _page_budget(base_size)
    steps = [
        lambda r: r.pop("activities", None) or r,
        lambda r: r.pop("awards", None) or r,
        lambda r: _drop_coursework(r),
        lambda r: _truncate_bullets(r, "experience", 4),
        lambda r: _truncate_bullets(r, "projects", 2),
        lambda r: _truncate_bullets(r, "experience", 3),
        lambda r: r.pop("certifications", None) or r,
        lambda r: _keep_top_projects(r, 2),       # drop only 3rd+ projects
        lambda r: _truncate_bullets(r, "projects", 1),
        lambda r: _truncate_bullets(r, "experience", 2),
        lambda r: _drop_last(r, "skills"),
        lambda r: _keep_top_projects(r, 1),        # last resort: single project
    ]
    for step in steps:
        if _estimate_line_count(r, base_size=base_size) <= budget:
            break
        step(r)
    return r


# ---------------------------------------------------------------------
# Expand helpers (used when the resume looks half-empty)
# ---------------------------------------------------------------------
def _master_lookup_experience(master: dict, role: str, company: str) -> dict | None:
    role_l = (role or "").lower()
    co_l = (company or "").lower()
    for e in master.get("experience", []) or []:
        head = (e.get("role", "") or "").lower().split("(")[0].strip()
        if (e.get("company", "") or "").lower() == co_l and (
            head in role_l or role_l.startswith(head[:18])
        ):
            return e
    return None


def _master_project_by_name(master: dict, name: str) -> dict | None:
    name_l = (name or "").lower()
    for p in master.get("projects", []) or []:
        if (p.get("name", "") or "").lower() == name_l:
            return p
    return None


def _clean_band_variants(variants, k: int, *, exclude: set[str] | None = None) -> list[str]:
    """Pick up to k master `bullets_all` variants that already render cleanly —
    a full single line or a full double — preferring single (densest), then the
    shortest doubles. Forbidden/overlong variants are skipped: expansion runs
    AFTER the LLM line-fit (in resume_builder, no LLM available), so a raw
    forbidden master bullet pulled in here would orphan-wrap on the page."""
    from .line_fitter import classify as _band
    exclude = exclude or set()
    singles, doubles = [], []
    for v in variants or []:
        if not isinstance(v, str) or v.lower()[:60] in exclude:
            continue
        b = _band(len(v))
        if b == "single":
            singles.append(v)
        elif b == "double":
            doubles.append(v)
    singles.sort(key=len, reverse=True)   # fullest single first
    doubles.sort(key=len)                 # shortest clean double first
    return (singles + doubles)[:k]


def _project_dups_experience(name: str, resume: dict) -> bool:
    """True if a project name's distinctive words already appear in an
    experience role or its lead bullets (so it'd render as a duplicate)."""
    import re as _re
    toks = set(_re.findall(r"[a-z0-9]+", (name or "").lower()))
    # Ignore generic words so "AI Resume Builder" doesn't false-match on "ai".
    toks -= {"a", "an", "the", "ai", "app", "web", "tool", "platform", "system", "and", "of"}
    if not toks:
        return False
    for e in resume.get("experience", []) or []:
        hay = (e.get("role", "") or "") + " " + " ".join((e.get("bullets") or [])[:2])
        haytoks = set(_re.findall(r"[a-z0-9]+", hay.lower()))
        if len(toks & haytoks) / len(toks) >= 0.6:
            return True
    return False


def _expand(r: dict, master: dict, base_size: float = _CAL_FONT) -> dict:
    """Pull additional content from master until estimate >= PAGE_LINE_MINIMUM
    (without exceeding PAGE_LINE_BUDGET).

    Strategy, in priority order:
      1. Top up experience bullets (most ATS-valuable space)
      2. Add a 2nd education entry with coursework
      3. Add a 3rd project from master
      4. Add coursework to existing education entry
      5. Add certifications row
      6. Add awards row
    """
    # Font-aware budgets shadow the module constants so the body below reads
    # naturally while staying correct for the configured font size.
    PAGE_LINE_MINIMUM = _page_minimum(base_size)
    PAGE_LINE_BUDGET = _page_budget(base_size)

    def lines() -> int:
        return _estimate_line_count(r, base_size=base_size)

    # 1. top up experience bullets to 4 each
    for entry in r.get("experience", []) or []:
        if lines() >= PAGE_LINE_MINIMUM:
            return r
        m = _master_lookup_experience(master, entry.get("role", ""), entry.get("company", ""))
        if not m:
            continue
        existing = entry.get("bullets") or []
        existing_norm = [b.lower()[:60] for b in existing]
        # Only top up with master variants that already render cleanly (we can't
        # line-fit them here — no LLM in the renderer).
        candidates = _clean_band_variants(
            m.get("bullets_all"), 5, exclude=set(existing_norm))
        for cand in candidates:
            if lines() >= PAGE_LINE_BUDGET - 2:
                break
            if cand.lower()[:60] in existing_norm:
                continue
            if len(existing) >= 5:
                break
            existing.append(cand)
            existing_norm.append(cand.lower()[:60])
        entry["bullets"] = existing[:5]

    if lines() >= PAGE_LINE_MINIMUM:
        return r

    # 2. add 2nd education entry from master if missing
    edu = list(r.get("education") or [])
    if len(edu) < 2 and master.get("education"):
        existing_schools = {(e.get("school", "") or "").lower() for e in edu}
        for me in master["education"]:
            if (me.get("school", "") or "").lower() in existing_schools:
                continue
            new_e = {
                "school": me.get("school", ""),
                "degree": me.get("degree", ""),
                "location": me.get("location", ""),
                "dates": f"{me.get('start_date','')} – {me.get('end_date','')}".strip(" –"),
                "gpa": me.get("gpa", ""),
            }
            cw = me.get("coursework")
            if isinstance(cw, list) and cw:
                new_e["coursework"] = ", ".join(cw[:6])
            elif isinstance(cw, str) and cw:
                new_e["coursework"] = cw
            edu.append(new_e)
            break
        r["education"] = edu

    if lines() >= PAGE_LINE_MINIMUM:
        return r

    # 3. add a 3rd project from master if not already present AND it doesn't
    # duplicate an experience entry (a flagship project may also be an experience
    # headline — the tailor's project/experience dedup runs before render, so we
    # must re-check here or expansion silently re-introduces the duplicate).
    projs = list(r.get("projects") or [])
    if len(projs) < 3 and master.get("projects"):
        present = {(p.get("name", "") or "").lower() for p in projs}
        for mp in master["projects"]:
            if (mp.get("name", "") or "").lower() in present:
                continue
            if _project_dups_experience(mp.get("name", ""), r):
                continue
            # Pull clean-band variants only (no LLM here to line-fit them).
            clean = _clean_band_variants(mp.get("bullets_all"), 2)
            if not clean:
                continue   # no cleanly-rendering bullets -> skip this project
            projs.append({
                "name": mp.get("name", ""),
                "tech": mp.get("tech", ""),
                "link": mp.get("link", ""),
                "bullets": clean,
            })
            break
        r["projects"] = projs[:3]

    if lines() >= PAGE_LINE_MINIMUM:
        return r

    # 4. add coursework to UW entry from master if missing
    for e in r.get("education", []) or []:
        if e.get("coursework"):
            continue
        m_edu = next(
            (m for m in master.get("education", []) or []
             if (m.get("school", "") or "").lower() == (e.get("school", "") or "").lower()),
            None,
        )
        if m_edu and m_edu.get("coursework"):
            cw = m_edu["coursework"]
            e["coursework"] = ", ".join(cw[:6]) if isinstance(cw, list) else str(cw)

    if lines() >= PAGE_LINE_MINIMUM:
        return r

    # 5. awards
    if not r.get("awards") and master.get("awards"):
        r["awards"] = list(master["awards"])[:3]

    if lines() >= PAGE_LINE_MINIMUM:
        return r

    # 6. activities (last expansion choice — niceties only when there's room)
    if not r.get("activities") and master.get("activities"):
        # Each activity is one line; pull only as many as fit the budget.
        budget_room = PAGE_LINE_BUDGET - lines() - 2  # HDR + safety cushion
        if budget_room >= 2:
            r["activities"] = list(master["activities"])[:int(budget_room)]

    # Pull through education extras (minor / specializations / honors) from master
    # since these come for "free" — they're already on the existing edu line block.
    for e in r.get("education", []) or []:
        m_edu = next(
            (m for m in master.get("education", []) or []
             if (m.get("school", "") or "").lower() == (e.get("school", "") or "").lower()),
            None,
        )
        if not m_edu:
            continue
        if not e.get("minor") and m_edu.get("minor"):
            e["minor"] = m_edu["minor"]
        if not e.get("specializations") and m_edu.get("specializations"):
            e["specializations"] = list(m_edu["specializations"])
        if not e.get("honors") and m_edu.get("honors"):
            e["honors"] = list(m_edu["honors"]) if isinstance(m_edu["honors"], list) else m_edu["honors"]

    return r


def _inject_certifications(r: dict, master: dict, base_size: float = _CAL_FONT) -> dict:
    """Add certifications if they are not present and budget allows.

    Certifications are always relevant (they're real credentials) so we inject
    them unconditionally after page-fit, rather than only during expansion.
    """
    if r.get("certifications") or not master.get("certifications"):
        return r
    est = _estimate_line_count(r, base_size=base_size)
    # HDR + 1 content line = ~2.6 "lines"; leave a 3-line cushion.
    if est <= _page_budget(base_size) - 3:
        r["certifications"] = list(master["certifications"])[:3]
    return r


def _fit_to_page(resume: dict, master: dict | None = None,
                 base_size: float = _CAL_FONT) -> dict:
    """Two-way page fit: shrink if overflow, expand if undersized.

    Always normalizes skills first so the line estimator sees the real shape.
    Certifications are injected as a guaranteed post-fit step when budget allows.
    """
    r = deepcopy(resume)
    r["skills"] = _normalize_skills(r.get("skills", []))

    budget = _page_budget(base_size)
    minimum = _page_minimum(base_size)
    est = _estimate_line_count(r, base_size=base_size)
    if est > budget:
        r = _shrink(r, base_size)
    elif master and est < minimum:
        r = _expand(r, master, base_size)
        if _estimate_line_count(r, base_size=base_size) > budget:
            r = _shrink(r, base_size)

    # Always try to add certifications after fit (they don't make the resume
    # look sparse — they fill space purposefully).
    if master:
        r = _inject_certifications(r, master, base_size)

    return r


def build_resume_onepage(resume: dict, user: dict, out_path: Path,
                         master: dict | None = None, **kwargs):
    """Render the resume with two-way page fit.

    Pass `master` to enable expansion when the LLM's tailored output is too
    short to fill the page (the most common failure mode in practice).
    """
    base_size = float(kwargs.get("base_size", 9.0))
    fitted = _fit_to_page(resume, master=master, base_size=base_size)
    build_resume_docx(fitted, user, out_path, **kwargs)
    est = _estimate_line_count(fitted, base_size=base_size)
    LOG.info("page-fit: %d lines @ %.0fpt (target %d–%d)",
             est, base_size, _page_minimum(base_size), _page_budget(base_size))
    return fitted
